from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"D:\luyuntiantong-main\luyuntiantong")
TMP = ROOT / "tmp" / "submission_review"
TEMPLATE = TMP / "附件：补充评审资料：挑战赛作品设计文档模板（含团队成员在成果中的贡献） (1).docx"
OUTPUT = ROOT / "output" / "documents" / "路云天瞳-挑战赛作品设计文档-优化版.docx"
MEDIA = TMP / "current_media"

BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(31, 31, 31)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "DCE6F1"
LIGHT_GRAY = "F2F2F2"


def set_font(run, size=10.5, bold=False, color=DARK, name="宋体"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def clear_body(doc):
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.35)
    section.right_margin = Cm(2.35)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.first_line_indent = Pt(21)

    if "List Bullet" not in [s.name for s in doc.styles]:
        bullet = doc.styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)
        bullet.base_style = normal
        bullet.paragraph_format.left_indent = Pt(21)
        bullet.paragraph_format.first_line_indent = Pt(-10.5)

    for style_name, size, before, after in (
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 11.5, 8, 4),
    ):
        if style_name not in [s.name for s in doc.styles]:
            doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.first_line_indent = Pt(0)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("路云天瞳｜挑战赛作品设计文档")
    set_font(r, 8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("匿名评审材料｜线上展示版")
    set_font(r, 8, color=MUTED)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("路云天瞳")
    set_font(r, 28, True, BLUE, "黑体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("基于时空图网络的多智能体车路协同遮挡感知系统")
    set_font(r, 16, True, DARK, "黑体")

    if (MEDIA / "image1.png").exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(20)
        p.add_run().add_picture(str(MEDIA / "image1.png"), height=Cm(10.2))

    info = doc.add_table(rows=4, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    labels = [
        ("所在赛道与赛项", "A-ICV（以报名系统记录为准）"),
        ("作品形态", "Web 在线平台 + 算法流程 + 场景回放 + 演示视频"),
        ("在线演示", "https://luyuntiantong-1.onrender.com"),
        ("材料属性", "匿名评审版，不包含学校、团队及个人身份信息"),
    ]
    for row, (label, value) in zip(info.rows, labels):
        row.cells[0].width = Cm(4.0)
        row.cells[1].width = Cm(11.3)
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for i, text in enumerate((label, value)):
            p = row.cells[i].paragraphs[0]
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run(text), 10, i == 0, DARK)
    doc.add_page_break()


def add_heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_para(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_font(r, 10.5, True)
        r = p.add_run(text[len(bold_lead):])
        set_font(r)
    else:
        set_font(p.add_run(text))
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.space_after = Pt(3)
        set_font(p.add_run(item))


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_margins(cell, 130, 160, 130, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = Pt(0)
    set_font(p.add_run(f"{title}："), 10.5, True, BLUE)
    set_font(p.add_run(text), 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc, path, caption, width=15.5):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Cm(width))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.first_line_indent = Pt(0)
    cp.paragraph_format.space_after = Pt(8)
    set_font(cp.add_run(caption), 9, color=MUTED)


def add_capability_table(doc):
    rows = [
        ("时序数据标准化", "统一轨迹字段、置信度过滤、重复帧去除、短缺失插值、长缺失截断、10Hz 重采样", "已实现并有自动化测试"),
        ("轨迹预测", "TorchScript 推理接口、批处理、热更新、耗时统计；在线轻量版无模型时显式常速度降级", "工程链路完成，真实模型依赖权重"),
        ("结果闭环", "预测结果、推理日志、异常记录写入 SQLite；WebSocket prediction topic 10Hz 推送", "已实现"),
        ("三维可视化", "历史轨迹实线、预测轨迹虚线、置信度红绿映射、16 类场景回放", "已实现"),
        ("误差指标", "ADE、FDE、Miss Rate；前端最近 100 个可对齐样本滑动统计", "真实数据不足时显示 --"),
        ("健康监控", "模型、推理耗时、GPU、SQLite、在线演示状态、预测日志查询", "已实现"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ["能力模块", "实现内容", "评审展示口径"]
    widths = [Cm(3.3), Cm(8.2), Cm(4.0)]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        set_font(p.add_run(h), 9.5, True)
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].width = widths[i]
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.first_line_indent = Pt(0)
            set_font(p.add_run(value), 9)


def add_contribution_table(doc):
    members = [
        ("队员1", "本科", "2024.09", "总体架构、STGNN 数据与训练流程、预测服务、风险决策、系统集成"),
        ("队员2", "本科", "2024.09", "Three.js 三维大屏、轨迹绘制、场景交互、展示模式与前端联调"),
        ("队员3", "本科", "2025.09", "YOLO/DeepSORT 感知流程、轨迹接口、场景数据清洗与测试"),
        ("队员4", "本科", "2024.09", "MQTT 通信、传感器/车端接口设计、硬件迁移方案与稳定性验证"),
        ("队员5", "本科", "2024.09", "FastAPI、SQLite 场景库、回放引擎、日志与 WebSocket 服务"),
        ("队员6", "本科", "2025.09", "项目统筹、匿名评审材料、演示脚本、视频与成果归档"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ["成员", "教育层次", "入学年月", "在作品中的主要分工"]
    widths = [Cm(2.0), Cm(2.5), Cm(2.5), Cm(8.7)]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        set_font(p.add_run(h), 9.5, True)
    for values in members:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].width = widths[i]
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.first_line_indent = Pt(0)
            if i < 3:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p.add_run(value), 9)

    doc.add_paragraph()
    add_para(doc, "主要成果及贡献口径：当前可核验成果包括作品源代码、在线演示平台、作品设计文档、场景库与自动化测试、作品介绍视频和汇报答辩视频。若截至提交时没有正式发表论文、授权专利或已登记软件著作权，应在成果排名栏统一填写“--”，不得用“计划申报”替代已取得成果。")


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(TEMPLATE)
    clear_body(doc)
    style_document(doc)
    add_title(doc)

    add_heading(doc, "一、目标问题与意义价值", 1)
    add_para(doc, "城市路口中，大型车辆、建筑物、绿化带和道路设施会形成视线遮挡。当行人、非机动车或车辆从遮挡区域进入冲突路径时，单车摄像头往往只能在目标露出后发现风险，留给驾驶员或自动驾驶系统的反应时间有限。作品面向这一类“鬼探头”与交叉冲突场景，利用路侧视角补足车端视野，并通过连续轨迹预测提升风险判断的提前量。")
    add_para(doc, "作品目标不是宣称视觉能够穿透障碍物，而是把路侧可见轨迹、目标交互和车辆状态转换为统一时序数据，在目标短时不可见或即将进入冲突区域时推演未来运动趋势，并形成“感知—预测—决策—记录—展示”的闭环。")
    add_bullets(doc, [
        "面向行人、非机动车和机动车等交通参与者，建立统一的 10Hz 轨迹数据口径。",
        "对未来约 2 秒轨迹进行预测，并结合 TTC、碰撞概率和制动建议辅助风险决策。",
        "将场景、预测、异常和推理日志持久化，支持实时观察、历史回放与指标评估。",
        "以低成本视觉设备和标准协议为主，保留向 Jetson、Atlas、摄像头和车载终端迁移的接口。",
    ])
    add_callout(doc, "线上提交边界", "本次作品提交为线上平台展示，仅支持 Web 界面、算法流程及演示视频的上传。硬件实物的实时交互、功耗、温度、端侧帧率和物理制动效果无法通过平台完整呈现，相关指标必须以线下实机记录为准。")

    add_heading(doc, "二、设计思路与方案", 1)
    add_heading(doc, "2.1 总体思路", 2)
    add_para(doc, "系统采用“路侧感知—云端预测—车端决策—Web 展示”的分层方案。路侧节点负责视频输入、目标检测、跨帧跟踪与道路坐标转换；云端维护目标历史轨迹，执行 STGNN 或显式降级预测，并将结果写入数据库与实时推送队列；车端依据共享感知和自身状态计算风险；Web 平台完成三维可视化、事件回放、健康监控和参数展示。")
    add_figure(doc, MEDIA / "image2.png", "图 1  系统总体架构与数据闭环", 14.0)

    add_heading(doc, "2.2 统一时序数据与轨迹样本", 2)
    add_para(doc, "每帧观测统一为 {track_id, class, x, y, vx, vy, timestamp, confidence}。数据进入训练或推理前依次完成低置信度过滤、重复帧去除、坐标边界校验、时间戳对齐和缺失补全：confidence < 0.3 的点被剔除；相同 track_id 与 timestamp 保留置信度最高记录；连续丢失不超过 3 帧时线性插值，超过 3 帧则截断轨迹；绝对坐标超过 200m 或非有限数值的点标记异常并丢弃；最终重采样到 10Hz。")
    add_para(doc, "TrajectoryDataset 支持 JSON、JSONL 和 SQLite 加载，训练样本默认使用过去 20 帧预测未来 20 帧，使离线训练、独立评估和在线推理共享同一字段口径。")

    add_heading(doc, "2.3 预测、监控与降级策略", 2)
    add_para(doc, "训练流程基于 PyTorch DataLoader 与 Adam 优化器，支持平移、旋转和高斯噪声增强，输出 ADE、FDE 与 Miss Rate（FDE > 2m）。训练完成后可导出 TorchScript，由 InferenceEngine 启动时预加载，支持批量推理、耗时统计和模型热更新。")
    add_para(doc, "在线轻量容器为了控制体积与算力，不捆绑 PyTorch 模型权重。当 TorchScript 不存在或未加载时，系统不会伪装成模型推理，而是显式返回 constant-velocity fallback，并在健康接口与算法面板中标注“常速度降级演示”。真实 STGNN 指标只有在加载模型并获得可对齐真值后才显示。")
    add_figure(doc, MEDIA / "image33.png", "图 2  ST-GNN 轨迹预测模型结构示意", 15.5)

    add_heading(doc, "三、方案实现", 1)
    add_heading(doc, "3.1 软件实现", 2)
    add_capability_table(doc)
    add_heading(doc, "3.2 数据库与接口", 2)
    add_bullets(doc, [
        "SQLite：frames、events、scenario_runs、predictions、inference_log、prediction_anomalies、experiments 等表。",
        "REST：/health、/api/v1/scenarios、/api/v1/demo/*、/api/v1/logs/prediction、/api/v1/prediction/reload。",
        "WebSocket：/api/v1/realtime/ws，支持 perception、prediction、decision、event、vehicle_status 等 topic 订阅。",
        "部署：Docker 单容器同源提供 React、FastAPI、REST 与 WebSocket，避免跨域和 HTTPS/WSS 混合内容问题。",
    ])
    add_heading(doc, "3.3 硬件与迁移接口", 2)
    add_para(doc, "目标硬件链路由路侧摄像头、边缘计算板、无线通信模块和车载终端组成。当前线上版本以 PC/容器、SQLite 场景库和标准协议复现业务数据流，保留 Jetson Orin Nano、Atlas 200 DK、摄像头和 OBU 的配置及迁移接口。评审材料中应区分“线上已实现的软件能力”“已设计的硬件接入方式”和“需要线下实机验证的物理指标”。")

    add_heading(doc, "四、运行结果与应用效果", 1)
    add_heading(doc, "4.1 当前可核验运行结果", 2)
    add_bullets(doc, [
        "在线演示地址可直接访问，首页、核心三维大屏与 /health 健康接口均由同一 HTTPS 域名提供。",
        "系统内置 16 类 GP/NM/IC 场景，支持启动、停止、单步、循环和确定性回放；默认在线场景以 10Hz 运行。",
        "实时消息带有 scenario_id、run_id 和 frame_id，可驱动三维目标、风险状态、历史轨迹和未来轨迹展示。",
        "模型未加载时明确展示降级状态；GPU、ADE、FDE、硬件温度等没有真实来源的指标显示为“--”，不使用随机值冒充实测。",
        "项目已通过数据集、预测运行时、部署契约、场景回放、前端构建等自动化验证。",
    ])
    add_figure(doc, MEDIA / "demo-night-intent.png", "图 3  夜间红外鬼探头场景：单车聚焦与协同驾驶意图", 15.5)
    add_figure(doc, MEDIA / "demo-night-flow.png", "图 4  夜间场景：全路网流量态势与车道状态", 15.5)
    add_figure(doc, MEDIA / "demo-pedestrian.png", "图 5  多行人连续穿越场景：俯视态势与目标跟踪", 15.5)
    add_figure(doc, MEDIA / "demo-yellow-light.png", "图 6  黄灯变红抢行场景：风险事件与协同制动意图", 15.5)

    add_heading(doc, "4.2 指标使用原则", 2)
    add_para(doc, "设计文档中的运行结果只采用可由代码、自动化测试、在线接口、训练日志或线下测试记录复现的数据。原稿中诸如特定召回率、MOTA、FDE、全链路 42ms、单路口成本降低 80% 等数字，如果没有对应数据集划分、实验脚本输出、测试设备、样本量和原始记录，不应作为既成结论提交。可在完成专项测试后，以“测试条件—数据规模—比较基线—指标结果”的完整格式补入。")
    add_callout(doc, "评审建议", "现场演示优先展示可复现闭环：选择场景 → 10Hz 实时播放 → 目标与轨迹更新 → 风险变化 → 事件生成 → 回放与日志查询。算法性能数字应另附实验记录或截图作为证据。")

    add_heading(doc, "五、创新与特色", 1)
    add_heading(doc, "5.1 面向遮挡风险的时空推理闭环", 2)
    add_para(doc, "作品将路侧观察、轨迹预测和车辆风险决策组织为统一时间线，不仅显示“当前目标在哪里”，还输出未来轨迹、风险等级和事件证据。")
    add_heading(doc, "5.2 训练与在线推理的数据口径一致", 2)
    add_para(doc, "TrajectoryDataset 将清洗、插值、截断和 10Hz 重采样集中实现，减少训练、评估与线上服务之间的字段漂移。")
    add_heading(doc, "5.3 可解释的工程降级", 2)
    add_para(doc, "模型、GPU或硬件遥测不可用时，系统显示原因和降级状态，而不是生成看似真实的随机指标；这增强了比赛展示的可信度和工程可维护性。")
    add_heading(doc, "5.4 场景库驱动的可复现演示", 2)
    add_para(doc, "16 类遮挡、非机动车和交叉冲突场景使用统一 scenario_id/run_id/frame_id 管理，可用于在线演示、协议回归、前端联调和算法评估。")
    add_heading(doc, "5.5 单容器在线交付", 2)
    add_para(doc, "React、FastAPI、SQLite 与 WebSocket 同源部署，评委无需安装开发环境即可访问完整作品流程，并保留平台自动域名作为比赛备用链接。")

    add_heading(doc, "六、团队成员在成果中的贡献", 1)
    add_contribution_table(doc)

    add_heading(doc, "附录 A：复现与验收入口", 1)
    add_bullets(doc, [
        "源代码：按组委会要求单独提交匿名 ZIP 包，设计文档中不展示可能关联参赛者身份的仓库账号。",
        "在线平台：https://luyuntiantong-1.onrender.com",
        "核心大屏：https://luyuntiantong-1.onrender.com/zhiluwujie",
        "健康接口：https://luyuntiantong-1.onrender.com/health",
        "建议评审浏览器：最新版 Chrome 或 Edge，分辨率 1920×1080。",
    ])
    add_para(doc, "说明：在线演示受云平台实例休眠、网络质量和免费资源额度影响，比赛当天应提前唤醒，并准备本地录屏与平台自动域名作为备用。")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
